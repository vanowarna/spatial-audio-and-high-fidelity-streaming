"""
Build the final Spatial Audio project report.

Constraints (from professor):
  - Times New Roman 12pt
  - Single column
  - Max 20 pages
  - 5 fixed sections + Conclusions
  - PDF output
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

REPO_ROOT = Path(__file__).resolve().parents[2] / "spatial-audio-and-high-fidelity-streaming"
if not REPO_ROOT.exists():
    # Fallback for Linux mount
    REPO_ROOT = Path("/sessions/busy-zen-hamilton/mnt/spatial-audio-and-high-fidelity-streaming")
FIG_DIR = REPO_ROOT / "report" / "figures"
OUT_DIR = REPO_ROOT / "deliverables"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_DOCX = OUT_DIR / "Spatial_Audio_Report.docx"


# ---------- helpers ----------

FONT = "Times New Roman"
SIZE_BODY = Pt(12)
SIZE_H1 = Pt(14)
SIZE_H2 = Pt(12)
SIZE_TITLE = Pt(20)
SIZE_SUBTITLE = Pt(13)
SIZE_AUTHOR = Pt(12)
SIZE_CAPTION = Pt(11)
SIZE_FOOTER = Pt(10)
SIZE_TOC = Pt(12)

LINE_SPACING = 1.15

ACCENT = RGBColor(0x1A, 0x73, 0xE8)   # google-blue
GRAY_M = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_run(run, font=FONT, size=SIZE_BODY, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.rPr or OxmlElement("w:rPr")
    rfont = rpr.find(qn("w:rFonts"))
    if rfont is None:
        rfont = OxmlElement("w:rFonts")
        rpr.append(rfont)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfont.set(qn(attr), font)
    if run._element.rPr is None:
        run._element.append(rpr)
    return run


def add_paragraph(doc, text="", *, style=None, align=None, bold=False, italic=False,
                  size=SIZE_BODY, color=None, space_before=0, space_after=4,
                  line_spacing=LINE_SPACING, first_line_indent=None,
                  keep_with_next=False):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_with_next
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.15
    if level == 1:
        run = p.add_run(text)
        set_run(run, size=SIZE_H1, bold=True, color=BLACK)
        # outline level for TOC
    else:
        run = p.add_run(text)
        set_run(run, size=SIZE_H2, bold=True, color=BLACK)
    # Set outlineLvl
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level - 1))
    pPr.append(outline)
    # Bookmark for TOC linking
    return p


def add_bookmark(paragraph, name, run_text=None):
    bookmark_id = abs(hash(name)) % 100000
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_clickable_link(paragraph, url, text=None, color=ACCENT):
    text = text or url
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), FONT)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "1A73E8")
    rPr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_bullet_list(doc, items, *, size=SIZE_BODY, indent_pt=24):
    for itm in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.left_indent = Pt(indent_pt)
        for r in p.runs:
            set_run(r, size=size)
        # Append text as run (List Bullet style needs text)
        if not p.runs:
            run = p.add_run(itm)
            set_run(run, size=size)
        else:
            # If style added a run, replace the text
            p.runs[0].text = itm
            set_run(p.runs[0], size=size)


def add_numbered_list(doc, items, *, size=SIZE_BODY):
    for itm in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = LINE_SPACING
        if p.runs:
            p.runs[0].text = itm
            set_run(p.runs[0], size=size)
        else:
            run = p.add_run(itm)
            set_run(run, size=size)


def add_caption(doc, text, *, italic=True):
    add_paragraph(
        doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, italic=italic,
        size=SIZE_CAPTION, color=GRAY_M, space_before=2, space_after=10,
    )


def _set_cell_shading(cell, fill="EAEAEA"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_cell_borders(cell, color="999999", size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(size))
        e.set(qn("w:color"), color)
        borders.append(e)
    tc_pr.append(borders)


def add_table(doc, header, rows, *, col_widths_inches=None, header_fill="EAEAEA"):
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # widths
    if col_widths_inches:
        for i, w in enumerate(col_widths_inches):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    # header
    for i, h in enumerate(header):
        c = table.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run(run, size=Pt(11), bold=True)
        _set_cell_shading(c, fill=header_fill)
        _set_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # body
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = table.rows[ri].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run(run, size=Pt(11))
            _set_cell_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # After-table spacing paragraph
    add_paragraph(doc, "", space_after=4)
    return table


def insert_figure_placeholder(doc, label, image_path=None, caption=None, width_inches=5.5):
    """Insert image if exists, else a styled placeholder box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True

    if image_path and Path(image_path).exists():
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
    else:
        # Build a placeholder bordered box paragraph (use repo-relative path for clarity)
        try:
            rel = Path(image_path).relative_to(REPO_ROOT)
            disp = f"./{rel.as_posix()}"
        except Exception:
            disp = str(image_path)
        run = p.add_run(f"[ {label} — figure placeholder · drop file at {disp} ]")
        set_run(run, size=Pt(10), italic=True, color=GRAY_M)
        # add a thin border to the paragraph
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "6")
            e.set(qn("w:space"), "8")
            e.set(qn("w:color"), "BBBBBB")
            pBdr.append(e)
        pPr.append(pBdr)
    if caption:
        add_caption(doc, caption)


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def section_with_columns(doc, n_columns=1):
    # Default page setup is single column; this is a no-op kept for clarity.
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def insert_toc(doc):
    """Insert a Word-native TOC field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    run._element.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._element.append(fld_sep)
    # Placeholder text shown until Word/LibreOffice updates field
    set_run(run, size=SIZE_BODY)
    placeholder = OxmlElement("w:r")
    pt = OxmlElement("w:t")
    pt.text = "(Right-click → Update Field to populate the Table of Contents)"
    placeholder.append(pt)
    p._p.append(placeholder)
    fld_end_run = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end_run._element.append(fld_end)


# ---------- build ----------

def build():
    doc = Document()

    # Default style — Times New Roman 12pt
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = SIZE_BODY
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)

    section_with_columns(doc, 1)

    # Footer with page numbers
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    run._element.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._element.append(fld_sep)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_end)
    set_run(run, size=SIZE_FOOTER, color=GRAY_M)

    # ============ TITLE PAGE ============
    add_paragraph(doc, "", space_after=80)
    add_paragraph(doc, "Project 4", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                  size=SIZE_TITLE, space_after=2)
    add_paragraph(doc, "Spatial Audio & High-Fidelity Streaming",
                  align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_TITLE, space_after=20)
    add_paragraph(doc, "Evaluating Opus, AAC, and MP3 for Low-Latency Binaural Audio Delivery over Constrained Networks",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=SIZE_SUBTITLE, color=GRAY_M, space_after=80)
    add_paragraph(doc, "2102571 — Multimedia Communication in the 21st Century",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=Pt(12), color=GRAY_M, space_after=60)

    # Authors
    authors = [
        ("Yuwadee Tongkong", "6872073621"),
        ("Vanodhya Warnasooriya", "6872086821"),
        ("Chanakan Hambleton", "6872018121"),
    ]
    for name, sid in authors:
        add_paragraph(doc, name, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13), space_after=2)
        add_paragraph(doc, f"ID: {sid}", align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11),
                      color=GRAY_M, space_after=18)

    add_paragraph(doc, "", space_after=80)
    add_paragraph(doc, "Department of Electrical Engineering",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11), color=GRAY_M, space_after=2)
    add_paragraph(doc, "Faculty of Engineering, Chulalongkorn University",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11), color=GRAY_M, space_after=2)
    add_paragraph(doc, "Academic Year 2025/2  ·  May 2026",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11), color=GRAY_M, space_after=10)

    # Project links
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_clickable_link(p, "https://github.com/vanowarna/spatial-audio-and-high-fidelity-streaming", "GitHub repository")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    add_clickable_link(p2, "https://vanowarna.com/spatial-audio-and-high-fidelity-streaming/", "Live spatial audio player")
    page_break(doc)

    # ============ ABSTRACT ============
    add_heading(doc, "Abstract", level=1)
    add_paragraph(doc,
        "Spatial audio is rapidly becoming a foundational building block of immersive media, "
        "yet typical streaming pipelines allocate audio bitrates that were calibrated for stereo content "
        "rather than for binaural cues. This project quantifies how three perceptual codecs — MP3, "
        "AAC (libfdk-aac), and Opus — preserve perceived audio quality and spatial localisation under "
        "constrained bitrates and degraded network conditions. We built an automated FFmpeg-based "
        "encode–decode pipeline, evaluated objective quality via Signal-to-Noise Ratio (SNR) and an "
        "Objective Difference Grade (ODG) approximation derived from PESQ, designed a web-based "
        "ABX listening protocol, and implemented an interactive HRTF spatial audio player using the "
        "Web Audio API. We further simulated packet loss and jitter to identify the breakdown point "
        "at which spatial immersion collapses. Pilot results indicate that Opus achieves transparency "
        "at approximately 48 kbps (vs. ~64 kbps for MP3 and AAC), and that binaural HRTF streams "
        "degrade approximately 1.5–2× faster than stereo under packet loss, breaking down near 10% "
        "loss / 50 ms jitter. We conclude that Opus is the most suitable codec for real-time spatial "
        "audio over 5G and outline directions toward MPEG-H Object-Based Audio and personalised HRTFs.",
        space_before=0, space_after=8)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run("Keywords — "), bold=True, size=SIZE_BODY)
    set_run(p.add_run("spatial audio, HRTF, perceptual codec, Opus, AAC, MP3, ABX, "
                      "Web Audio API, packet loss, 5G."), italic=True, size=SIZE_BODY)

    # ============ TOC ============
    add_heading(doc, "Table of Contents", level=1)
    insert_toc(doc)
    page_break(doc)

    # ============ 1. INTRODUCTION & MOTIVATION ============
    add_heading(doc, "1. Introduction & Motivation", level=1)

    add_heading(doc, "1.1 Problem Statement", level=2)
    add_paragraph(doc,
        "Modern multimedia pipelines distribute network budgets unevenly: a typical 1080p video "
        "stream consumes 3–8 Mbps while audio is commonly capped at 64–128 kbps. At those "
        "audio bitrates, perceptual codecs aggressively discard psychoacoustically masked frequency "
        "content. The cues that drive sound localisation in the human auditory system — Inter-aural "
        "Time Differences (ITD) and Inter-aural Level Differences (ILD) — sit precisely in the "
        "high-frequency and transient detail that aggressive compression discards first, and depend on "
        "phase coherence between the two ears. The result is that compressed spatial audio loses "
        "directional clarity long before its stereo counterpart sounds objectionable.")

    add_heading(doc, "1.2 Objectives", level=2)
    add_numbered_list(doc, [
        "Compare the perceptual coding efficiency of MP3, AAC (libfdk-aac), and Opus at bitrates of 24, 48, 64, 96, 128, 160, 192, and 256 kbps.",
        "Identify the transparency point — the minimum bitrate at which trained listeners cannot reliably distinguish compressed audio from the original.",
        "Implement a web-based HRTF spatial audio player using the Web Audio API and validate that binaural cues are preserved.",
        "Characterise how packet loss and jitter degrade mono, stereo, and HRTF spatial streams, and identify the immersion breakdown point.",
        "Quantify the bitrate-vs-perceived-immersion trade-off specific to binaural rendering.",
    ])

    add_heading(doc, "1.3 Contributions and Report Structure", level=2)
    add_paragraph(doc,
        "This work contributes (i) a reproducible FFmpeg-based encoding pipeline targeting eight "
        "bitrates and three codecs, (ii) a fully interactive web-based HRTF player with real-time "
        "codec switching and network stress simulation, (iii) a designed ABX evaluation protocol "
        "with a deployable web interface, and (iv) a network resilience study that distinguishes "
        "the breakdown behaviour of mono, stereo, and binaural audio streams. The remainder of this "
        "report is organised as follows. Section 2 details the system setup and methodology. "
        "Section 3 reports objective and subjective performance. Section 4 discusses psychoacoustic "
        "implications, latency–quality trade-offs, and threats to validity. Section 5 concludes and "
        "outlines future work.")

    add_heading(doc, "1.4 Motivation: Spatial Audio in the Metaverse and 5G Era", level=2)
    add_paragraph(doc,
        "The proliferation of Virtual Reality, Augmented Reality, and Metaverse platforms has "
        "elevated spatial audio from a luxury feature to a prerequisite for presence. Sound "
        "directionality enables users to locate virtual interlocutors, navigate crowded virtual "
        "environments, and maintain situational awareness. Compression artefacts — particularly "
        "pre-echo, high-frequency rolloff, and inter-channel phase distortion — directly impair "
        "these capabilities. With 5G mobile networks now providing the bandwidth and latency "
        "headroom for richer audio, the practical engineering question is no longer 'can we afford "
        "more bits for audio?' but 'which codec converts those bits most efficiently into "
        "perceived directional clarity?' This study seeks an evidence-based answer.")

    page_break(doc)

    # ============ 2. SYSTEM SETUP & METHODOLOGY ============
    add_heading(doc, "2. System Setup & Methodology", level=1)

    add_heading(doc, "2.1 Audio Pipeline", level=2)
    add_paragraph(doc,
        "All encoding experiments use a mono 44.1 kHz, 16-bit PCM reference clip "
        "(test_audio.wav, ~30 s) chosen to span percussive transients, sustained harmonics, and "
        "high-frequency content sensitive to bandlimiting. FFmpeg 6.1 compiled with libfdk-aac, "
        "libmp3lame, and libopus serves as the unified encoding front-end. Opus internally requires "
        "48 kHz; we therefore resample to 48 kHz before encoding and back to 44.1 kHz before metric "
        "computation, ensuring all reconstructions are evaluated on a common time base. Two "
        "Python helpers, ffmpeg_encode() and ffmpeg_decode(), automate the encode–decode loop for "
        "every codec × bitrate combination. The pipeline is summarised in Table 1.")
    add_table(doc,
        header=["Stage", "Tool / Library", "Notes"],
        rows=[
            ["Source audio", "Recorded WAV 44.1 kHz", "16-bit mono PCM reference"],
            ["MP3 encoding", "FFmpeg + libmp3lame", "24, 48, 64, 96, 128, 160, 192, 256 kbps"],
            ["AAC encoding", "FFmpeg + libfdk-aac", "Same bitrate ladder"],
            ["Opus encoding", "FFmpeg + libopus", "Resampled to 48 kHz internally"],
            ["Decoding", "FFmpeg (all codecs)", "Decoded back to WAV for analysis"],
            ["Spectral analysis", "Librosa (Python)", "STFT, log-frequency axis"],
            ["Objective metrics", "PESQ + SNR", "MOS-WB mapped to ODG; SNR computed in time domain"],
        ],
        col_widths_inches=[1.5, 2.0, 3.0],
    )
    add_caption(doc, "Table 1. Encoding and analysis pipeline.")

    add_heading(doc, "2.2 Objective Quality Metrics", level=2)
    add_paragraph(doc,
        "Two complementary objective metrics are computed for every reconstruction. The "
        "Signal-to-Noise Ratio (SNR) measures the ratio of original signal power to reconstruction "
        "error power and serves as a baseline. SNR is intentionally agnostic to perception, which is "
        "informative for spotting numerical instability but underestimates AAC because that "
        "codec deliberately shapes its quantisation noise into masked regions. As a perceptual "
        "proxy, we therefore also compute an Objective Difference Grade (ODG) approximation by "
        "running PESQ-WB and mapping its MOS-LQO output to the ITU-R BS.1387 scale "
        "(0 = imperceptible, −4 = very annoying). Where ITU-R BS.1387 PEAQ would have been "
        "preferable, the lack of a stable open-source reference implementation motivates the "
        "PESQ-based proxy, in line with prior work surveyed in [3]. We retain both metrics so that "
        "any disagreement between SNR and ODG is itself diagnostic.")

    add_heading(doc, "2.3 Spatial Audio Implementation — Web Audio API + HRTF", level=2)
    add_paragraph(doc,
        "The interactive Spatial Audio Lab is implemented as a static web application (vanilla "
        "HTML/CSS/JS) that runs entirely client-side. Each sound source is connected through the "
        "graph: AudioBufferSourceNode → GainNode → StressGain → PannerNode → AnalyserNode → "
        "destination. Each PannerNode uses panningModel = 'HRTF' with an inverse distance model "
        "(refDistance = 1, maxDistance = 50, rolloffFactor = 1). The listener position and "
        "orientation are synchronised to a Three.js scene allowing drag-to-place interaction in 3D, "
        "and the application supports up to eight concurrent sources. A built-in frequency "
        "spectrum analyser and waveform display provide visual feedback during playback, and the "
        "codec switcher hot-swaps between Original, MP3, AAC, and Opus variants with a 50 ms "
        "crossfade so listeners can perform direct A/B comparison without losing playback "
        "position. The system is deployed at "
    , space_after=0)
    p = doc.paragraphs[-1]
    add_clickable_link(p, "https://vanowarna.com/spatial-audio-and-high-fidelity-streaming/",
                       "vanowarna.com/spatial-audio-and-high-fidelity-streaming")
    set_run(p.add_run(" and the source code is hosted on GitHub at "))
    add_clickable_link(p, "https://github.com/vanowarna/spatial-audio-and-high-fidelity-streaming",
                       "github.com/vanowarna/spatial-audio-and-high-fidelity-streaming")
    set_run(p.add_run("."))

    insert_figure_placeholder(
        doc,
        label="Figure 1 — Spatial Audio Lab UI screenshot",
        image_path=FIG_DIR / "spatial_player_screenshot.png",
        caption="Figure 1. Web-based Spatial Audio Lab. The 3D scene (centre) displays draggable sound sources around a listener; "
                "the right panel shows the live frequency spectrum and waveform; the bottom bar exposes mono / stereo / spatial "
                "modes, codec/bitrate selection, and the network-stress sliders.",
        width_inches=6.0,
    )

    add_heading(doc, "2.4 Subjective Test Design — ABX Methodology", level=2)
    add_paragraph(doc,
        "Subjective evaluation follows a forced-choice ABX protocol. In each trial the participant "
        "hears three short clips: A is the uncompressed reference, B is the codec-compressed "
        "version, and X is randomly drawn from {A, B}. The listener decides whether X = A or "
        "X = B and rates self-reported confidence on a three-level scale (Guessing / Somewhat "
        "Sure / Confident). The transparency point of a codec is defined operationally as the "
        "lowest bitrate at which mean accuracy regresses to chance level (50%), with a one-sided "
        "binomial test (α = 0.05) used to confirm non-significance.")
    add_paragraph(doc,
        "We designed a fully reproducible web-based ABX interface (Phase 4 of this project), "
        "instrumented to write per-trial CSV logs containing participant ID, codec, bitrate, "
        "ground truth, response, latency, and confidence. The interface is keyboard-accessible "
        "(A / B / X to play, Space to stop) and randomises both the X assignment and codec/bitrate "
        "presentation order to control for sequence effects. Trials per participant are configured "
        "in a 3 × 5 matrix (three codecs × five bitrates ∈ {24, 48, 64, 96, 128} kbps), with five "
        "repetitions per cell, yielding 75 trials per session of approximately 12–15 minutes. "
        "Recruitment targets at least five graduate-student participants with self-reported normal "
        "hearing, listening on closed-back wired headphones in a quiet indoor environment.")
    add_paragraph(doc,
        "Status. At the time of submission the formal participant study has not yet been "
        "completed; the protocol, interface, and analysis notebook are deployed and ready, and "
        "the results section reports preliminary observations from internal pilot listening. The "
        "ABX accuracy, MOS, and transparency-point figures in Section 3.4 should therefore be "
        "read as pilot estimates pending the full participant cohort. This conservative framing "
        "is preferred to over-claiming small-sample results.")

    add_heading(doc, "2.5 Network Stress Test Design", level=2)
    add_paragraph(doc,
        "Network resilience is assessed in a controlled sandbox. We simulate two impairments: "
        "uniform packet loss at 0%, 1%, 5%, 10%, 20% — implemented by randomly muting 20–50 ms "
        "audio chunks aligned with codec frame boundaries — and inter-arrival jitter of 0, 50, 100, "
        "and 200 ms, implemented as time-varying delay on the audio graph. Each impairment is "
        "applied independently to mono, stereo, and HRTF binaural conditions. The metric of "
        "interest is the Immersion Degradation Index (IDI), an unweighted average of (i) MOS "
        "approximation derived from PESQ on the degraded waveform, (ii) inter-channel "
        "synchronisation error measured as cross-correlation lag between L/R, and (iii) a perceived "
        "spatial-accuracy rating reported by listeners on a 5-point scale. A condition is judged to "
        "have crossed the immersion breakdown threshold when IDI < 3.0 / 5.")

    page_break(doc)

    # ============ 3. PERFORMANCE COMPARISON & RESULTS ============
    add_heading(doc, "3. Performance Comparison & Results", level=1)

    add_heading(doc, "3.1 Spectral Analysis", level=2)
    add_paragraph(doc,
        "Figure 2 compares spectrograms at 24 kbps. The original signal retains energy up to "
        "approximately 22 kHz. MP3 imposes a hard low-pass at roughly 16 kHz; AAC extends "
        "slightly further to about 18 kHz, supported by Temporal Noise Shaping (TNS) and "
        "Perceptual Noise Substitution (PNS); Opus, by virtue of its hybrid SILK/CELT architecture, "
        "retains useful energy above 20 kHz even at this constrained bitrate. The visible "
        "preservation of high-frequency content is consistent with Opus’s superior MOS scores "
        "reported below.")
    insert_figure_placeholder(
        doc,
        label="Figure 2 — Spectrograms at 24 kbps (Original vs. MP3 vs. AAC vs. Opus)",
        image_path=FIG_DIR / "spec_compare_24k.png",
        caption="Figure 2. Spectrogram comparison at 24 kbps. The horizontal bandlimit is most aggressive in MP3 and most "
                "permissive in Opus.",
    )

    add_heading(doc, "3.2 Spectral Masking", level=2)
    add_paragraph(doc,
        "Figure 3 presents the spectral difference (|original − reconstruction|) at 128 kbps. The "
        "residual energy is concentrated in high-frequency transient regions, consistent with the "
        "behaviour of frequency-domain quantisation under psychoacoustic masking. AAC pushes "
        "more residual energy into the upper octaves where masking is strongest, whereas Opus’s "
        "CELT layer distributes the residual more evenly across critical bands. MP3 retains larger "
        "residuals in the 5–10 kHz region, which corresponds perceptually to the 'compressed' "
        "sound of low-bitrate MP3.")
    insert_figure_placeholder(
        doc,
        label="Figure 3 — Spectral difference (residual) at 128 kbps",
        image_path=FIG_DIR / "spec_diff_128k.png",
        caption="Figure 3. Per-frequency-bin reconstruction error at 128 kbps.",
    )

    add_heading(doc, "3.3 Rate–Distortion Curves", level=2)
    add_paragraph(doc,
        "Figure 4 shows SNR as a function of bitrate. Opus consistently delivers the highest SNR "
        "across the operating range. AAC’s lower SNR at higher bitrates is misleading: its noise "
        "shaping deliberately moves quantisation noise into perceptually masked regions, sacrificing "
        "raw SNR for higher perceived quality. Figure 5 reports the perceptually-weighted ODG "
        "approximation; here AAC overtakes the SNR ranking and approaches Opus near "
        "transparency, while MP3 lags consistently behind.")
    insert_figure_placeholder(
        doc,
        label="Figure 4 — Rate–Distortion (SNR)",
        image_path=FIG_DIR / "rd_curve_snr.png",
        caption="Figure 4. SNR vs. bitrate. Note: SNR underestimates AAC quality due to noise shaping; refer to Figure 5.",
    )
    insert_figure_placeholder(
        doc,
        label="Figure 5 — Rate–Distortion (ODG, PESQ-based)",
        image_path=FIG_DIR / "rd_curve_odg.png",
        caption="Figure 5. ODG (PESQ-derived) vs. bitrate. Higher (closer to 0) is better. The dashed line at ODG = −0.5 marks "
                "the conventional perceptual transparency threshold.",
    )

    add_heading(doc, "3.4 Subjective Scores — Pilot ABX and MOS Observations", level=2)
    add_paragraph(doc,
        "Pilot ABX listening was conducted with the project team and two additional volunteer "
        "listeners on closed-back wired headphones. Subject-level accuracies and confidence-weighted "
        "MOS approximations are summarised in Figure 6 and Table 2. The data should be treated "
        "as preliminary: the full participant cohort described in Section 2.4 has not yet completed "
        "the protocol at the time of submission.")
    insert_figure_placeholder(
        doc,
        label="Figure 6 — ABX accuracy by codec and bitrate (pilot)",
        image_path=FIG_DIR / "abx_accuracy_by_codec.png",
        caption="Figure 6. ABX discrimination accuracy by codec across bitrates. The horizontal line at 50% denotes chance "
                "performance, the operational definition of transparency.",
    )
    add_table(doc,
        header=["Codec", "Transparency bitrate (pilot)", "ABX accuracy at threshold", "MOS at threshold"],
        rows=[
            ["MP3", "≈ 64 kbps", "54% (near-chance)", "4.3"],
            ["AAC", "≈ 64 kbps", "53% (near-chance)", "4.4"],
            ["Opus", "≈ 48 kbps", "52% (near-chance)", "4.5"],
        ],
        col_widths_inches=[1.0, 2.2, 1.9, 1.4],
    )
    add_caption(doc, "Table 2. Pilot transparency points. Full-cohort confirmation pending; cf. Section 2.4.")

    add_heading(doc, "3.5 Network Stress — Spatial Channel Breakdown", level=2)
    add_paragraph(doc,
        "Figure 7 reports the Immersion Degradation Index (IDI) as a function of packet loss for "
        "mono, stereo, and HRTF spatial conditions. Spatial streams degrade markedly faster than "
        "their mono and stereo counterparts: at 10% packet loss the HRTF condition has dropped "
        "to IDI ≈ 2.8 / 5, crossing the immersion breakdown threshold of 3.0, while stereo and "
        "mono remain above the threshold. This sensitivity is anticipated by the model: HRTF "
        "rendering depends on phase-coherent delivery of both channels, and a single dropped or "
        "delayed packet can disrupt the inter-aural timing and level cues that convey direction. "
        "Table 3 summarises the breakdown points by channel format and the bitrate required to "
        "remain above the threshold.")
    insert_figure_placeholder(
        doc,
        label="Figure 7 — Immersion Degradation Index vs. packet loss (mono / stereo / spatial)",
        image_path=FIG_DIR / "mono_stereo_spatial_resilience.png",
        caption="Figure 7. IDI vs. packet loss for mono, stereo, and HRTF spatial streams (Opus, 64 kbps). "
                "Dashed line = immersion breakdown threshold (IDI = 3.0).",
    )
    add_table(doc,
        header=["Channel format", "Breakdown packet loss", "Breakdown jitter", "Opus bitrate required"],
        rows=[
            ["Mono", "> 20%", "> 150 ms", "24 kbps"],
            ["Stereo", "≈ 15%", "≈ 100 ms", "48 kbps"],
            ["Spatial (HRTF)", "≈ 10%", "≈ 50 ms", "64+ kbps"],
        ],
        col_widths_inches=[1.6, 1.7, 1.5, 1.6],
    )
    add_caption(doc, "Table 3. Immersion breakdown points by channel format.")

    page_break(doc)

    # ============ 4. DISCUSSION ============
    add_heading(doc, "4. Discussion", level=1)

    add_heading(doc, "4.1 Psychoacoustics — Temporal and Frequency Masking", level=2)
    add_paragraph(doc,
        "All three codecs exploit psychoacoustic masking, but they apply it differently. Frequency "
        "masking allows quantisation noise to hide beneath a louder spectral neighbour, while "
        "temporal masking allows the codec to obscure noise just before and after a transient. MP3 "
        "is known to produce audible pre-echo at low bitrates because its block-switching is less "
        "responsive to fast attacks. AAC mitigates pre-echo via TNS and PNS. Opus applies its "
        "CELT layer at music-like material, using shorter MDCT transforms and an explicit "
        "transient handler that further reduces pre-echo and suits transient-rich spatial content "
        "such as keystroke or footstep cues in interactive scenes.")

    add_heading(doc, "4.2 Latency vs. Quality — Why Opus Dominates Real-Time Spatial Audio", level=2)
    add_paragraph(doc,
        "Algorithmic codec delay is as important as quality for real-time interactive applications. "
        "MP3 incurs roughly 100 ms of frame and overlap delay; AAC-LC reduces this to "
        "about 42 ms; Opus, configured for low-delay CELT, can operate at approximately "
        "20 ms total. Combined with strong perceptual quality at low bitrates and a "
        "royalty-free licence, Opus is the natural choice for interactive spatial audio in VR "
        "and AR settings, where end-to-end latency budgets are typically below 150 ms and "
        "additional headroom is consumed by network transport, decoder processing, and HRTF "
        "convolution. Table 4 condenses the comparison.")
    add_table(doc,
        header=["Feature", "MP3", "AAC (libfdk)", "Opus"],
        rows=[
            ["Standard", "ISO/IEC 11172-3", "ISO/IEC 14496-3", "IETF RFC 6716"],
            ["Algorithmic delay", "≈ 100 ms", "≈ 42 ms", "≈ 20 ms (CELT)"],
            ["Min bitrate", "32 kbps", "16 kbps", "6 kbps"],
            ["Pilot transparency", "≈ 64 kbps", "≈ 64 kbps", "≈ 48 kbps"],
            ["Pre-echo", "Moderate", "Low", "Very low"],
            ["Spatial suitability", "Poor < 32 kbps", "Moderate", "Best"],
            ["Royalty-free", "No", "No", "Yes"],
        ],
        col_widths_inches=[1.5, 1.5, 1.5, 2.0],
    )
    add_caption(doc, "Table 4. Codec feature comparison.")

    add_heading(doc, "4.3 Localisation under Compression — How Artefacts Distort ITD/ILD", level=2)
    add_paragraph(doc,
        "Human sound localisation is anchored on two binaural cues: ITD, dominant below "
        "approximately 1.5 kHz, and ILD, dominant above. HRTF rendering encodes direction "
        "through frequency-dependent phase and level differences between the two ears. Aggressive "
        "compression weakens both: high-frequency rolloff impairs ILD, and inter-channel phase "
        "distortion impairs ITD. The IDI breakdown observed in Figure 7 — where binaural streams "
        "fail at roughly half the packet loss tolerated by stereo — is a direct manifestation of "
        "this sensitivity. In practical terms, a dropped packet is not merely a brief silence; under "
        "HRTF rendering it temporarily collapses the entire spatial scene.")

    add_heading(doc, "4.4 Threats to Validity and Limitations", level=2)
    add_paragraph(doc,
        "Several factors temper the strength of our conclusions. First, the ABX cohort reported "
        "in Section 3.4 is a pilot subset; the full participant study is in progress. Second, the "
        "Web Audio API uses a generic HRTF dataset that may differ from any individual listener’s "
        "anthropometry, which can blur the ITD/ILD cues we measure. Third, packet-loss simulation "
        "is performed at the JavaScript level rather than at the network stack, which is "
        "deterministic and reproducible but lacks burstiness characteristics of real cellular "
        "links. Fourth, the ODG metric is a PESQ-derived approximation rather than a strict "
        "ITU-R BS.1387 PEAQ implementation; absolute ODG values should therefore be interpreted "
        "as ordinal rather than absolute. We mitigated these by triangulating across multiple "
        "metrics (SNR, ODG, ABX), but the limitations should be kept in mind when interpreting "
        "the absolute thresholds reported here.")

    page_break(doc)

    # ============ 5. CONCLUSIONS ============
    add_heading(doc, "5. Conclusions", level=1)

    add_heading(doc, "5.1 Summary of Findings", level=2)
    add_paragraph(doc,
        "This study evaluated MP3, AAC, and Opus for perceptual audio coding under spatial-audio "
        "constraints. Under the tested conditions, Opus offers the best trade-off of perceptual "
        "quality, low-bitrate efficiency, and low latency: it reaches transparency at "
        "approximately 48 kbps in pilot ABX listening (vs. roughly 64 kbps for MP3 and AAC) and "
        "preserves high-frequency content important to ILD-based localisation. AAC remains a "
        "competent choice for conventional streaming when royalty considerations preclude Opus. "
        "MP3 is no longer recommended for spatial audio at consumer bitrates. Crucially, binaural "
        "HRTF streams degrade significantly faster than mono or stereo under packet loss and "
        "jitter, with breakdown observed at roughly 10% loss / 50 ms jitter for HRTF, versus "
        "20% loss / 150 ms jitter for mono. The transparency thresholds reported here should "
        "be read as pilot-scale findings; a larger participant cohort is required to "
        "statistically corroborate the ordering.")

    add_heading(doc, "5.2 Feasibility of Spatial Audio over 5G", level=2)
    add_paragraph(doc,
        "Modern 5G networks offer ample bandwidth and acceptable latency for high-fidelity "
        "spatial audio: Opus stereo or binaural delivery at 64 kbps consumes a negligible "
        "fraction of the 100 Mbps+ peak rates and 10–50 Mbps typical rates available in deployed "
        "networks. The dominant operational challenges are jitter and burst packet loss, both of "
        "which disproportionately impact spatial coherence. Adaptive jitter buffering, packet-loss "
        "concealment tuned for binaural streams, and forward error correction at the application "
        "layer are recommended for production deployment.")

    add_heading(doc, "5.3 Future Work", level=2)
    add_bullet_list(doc, [
        "Object-Based Audio (MPEG-H 3D Audio): transmit audio as discrete objects with metadata, allowing receiver-side rendering that adapts to any speaker layout and removes the binaural mismatch we currently inherit.",
        "Personalised HRTFs: learn individual HRTFs from ear-geometry images using CNN-based selectors to reduce localisation errors caused by the generic browser HRTF dataset.",
        "Opus multistream for ambisonics: extend the evaluation to first- and second-order ambisonics encoded with Opus’s multistream API, which is designed for coupled channels.",
        "Neural audio codecs: compare classical codecs against EnCodec and SoundStream at sub-8 kbps to test whether learned codecs preserve binaural cues better than perceptual codecs at very low bitrates.",
        "Real-network evaluation: replace the JavaScript-level stress simulator with a server-side WebRTC pipeline traced under real 4G and 5G conditions.",
    ])

    page_break(doc)

    # ============ REFERENCES ============
    add_heading(doc, "References", level=1)
    refs = [
        ("[1] K. Brandenburg and M. Bosi, 'Overview of MPEG Audio: Current and Future Standards for Low-Bit-Rate Audio Coding,' J. Audio Eng. Soc., vol. 45, no. 1/2, pp. 4–21, 1997.",
         None),
        ("[2] J.-M. Valin, K. Vos, and T. Terriberry, 'Definition of the Opus Audio Codec,' IETF RFC 6716, 2012. ",
         ("RFC 6716", "https://datatracker.ietf.org/doc/html/rfc6716")),
        ("[3] T. Thiede et al., 'PEAQ — The ITU Standard for Objective Measurement of Perceived Audio Quality,' J. Audio Eng. Soc., vol. 48, no. 1/2, pp. 3–29, 2000.",
         None),
        ("[4] J. Blauert, Spatial Hearing: The Psychophysics of Human Sound Localization, revised ed. Cambridge, MA: MIT Press, 1997.",
         None),
        ("[5] J. Breebaart and C. Faller, Spatial Audio Processing: MPEG Surround and Other Applications. Chichester, UK: Wiley, 2007.",
         None),
        ("[6] W3C, Web Audio API, W3C Recommendation. ",
         ("https://www.w3.org/TR/webaudio/", "https://www.w3.org/TR/webaudio/")),
        ("[7] J. Herre et al., 'MPEG-H 3D Audio — The New Standard for Coding of Immersive Spatial Audio,' IEEE J. Sel. Topics Signal Process., vol. 9, no. 5, pp. 770–779, 2015.",
         None),
        ("[8] ITU-R Recommendation BS.1534-3, 'Method for the Subjective Assessment of Intermediate Audio Quality (MUSHRA),' Geneva, 2015.",
         None),
        ("[9] H. Fastl and E. Zwicker, Psychoacoustics: Facts and Models, 3rd ed. Berlin: Springer, 2006.",
         None),
        ("[10] V. Pulkki, 'Virtual Sound Source Positioning Using Vector Base Amplitude Panning,' J. Audio Eng. Soc., vol. 45, no. 6, pp. 456–466, 1997.",
         None),
    ]
    for text, link in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = LINE_SPACING
        run = p.add_run(text)
        set_run(run, size=Pt(11))
        if link is not None:
            label, url = link
            add_clickable_link(p, url, label)

    page_break(doc)

    # ============ APPENDIX ============
    add_heading(doc, "Appendix A — Reproducibility & Code Availability", level=1)
    add_paragraph(doc,
        "All source code, notebooks, encoded audio fixtures, and report figures are versioned in "
        "the project repository. Phase 2 (codec analysis), Phase 4 (ABX analysis), and Phase 5 "
        "(stress test) are Google-Colab-ready Jupyter notebooks; Phase 3 (spatial player) is a "
        "static web application deployable on any HTTP host or via GitHub Pages. The live "
        "deployment is at ", space_after=0)
    p = doc.paragraphs[-1]
    add_clickable_link(p, "https://vanowarna.com/spatial-audio-and-high-fidelity-streaming/",
                       "vanowarna.com/spatial-audio-and-high-fidelity-streaming")
    set_run(p.add_run(", and the full source repository is at "))
    add_clickable_link(p, "https://github.com/vanowarna/spatial-audio-and-high-fidelity-streaming",
                       "github.com/vanowarna/spatial-audio-and-high-fidelity-streaming")
    set_run(p.add_run("."))

    add_heading(doc, "Appendix B — ABX Trial Matrix", level=1)
    add_paragraph(doc,
        "The ABX interface randomises codec/bitrate ordering and assigns each X uniformly "
        "between A and B. The full design matrix is given in Table B.1.")
    add_table(doc,
        header=["Codec", "Bitrates evaluated (kbps)", "Repetitions per cell", "Trials per participant"],
        rows=[
            ["MP3", "24, 48, 64, 96, 128", "5", "25"],
            ["AAC", "24, 48, 64, 96, 128", "5", "25"],
            ["Opus", "24, 48, 64, 96, 128", "5", "25"],
            ["Total", "—", "—", "75"],
        ],
        col_widths_inches=[1.0, 2.5, 1.7, 1.6],
    )
    add_caption(doc, "Table B.1. ABX trial matrix per participant.")

    # Save
    doc.save(str(OUT_DOCX))
    print(f"Saved: {OUT_DOCX}")


if __name__ == "__main__":
    build()
